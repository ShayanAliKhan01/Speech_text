import streamlit as st
import speech_recognition as sr
from deep_translator import GoogleTranslator
import docx
import io
import time

# Set page config
st.set_page_config(page_title="Speech-to-Text Translator", page_icon="🎤", layout="wide")

# Initialize session state variables
if 'recognized_text' not in st.session_state:
    st.session_state.recognized_text = ""
if 'translated_text' not in st.session_state:
    st.session_state.translated_text = ""
if 'target_language' not in st.session_state:
    st.session_state.target_language = "ur"
if 'translation_history' not in st.session_state:
    st.session_state.translation_history = []

# Language options
language_options = {
    "French": "fr",
    "Spanish": "es",
    "German": "de",
    "Urdu": "ur",
    "Hindi": "hi",
    "Arabic": "ar"
}

# Sidebar settings
st.sidebar.header("🌍 Language")
selected_language = st.sidebar.selectbox(
    "Select Translation Language:",
    options=list(language_options.keys()),
    index=list(language_options.values()).index(st.session_state.target_language) if st.session_state.target_language in list(language_options.values()) else 0
)
st.session_state.target_language = language_options[selected_language]

# Main Content (Vertical)
st.title("🎤 Speech-to-Text Translator")
st.write("Convert speech into text and translate it into different languages.")

st.divider()

### 🎙 Speech Recognition Section
st.subheader("🎙 Speech Recognition")

def recognize_speech():
    """Continuously listens for speech until 'stop recording' is detected."""
    recognizer = sr.Recognizer()
    recognizer.energy_threshold = 300  # Adjusts background noise sensitivity

    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source, duration=1)
        st.info("Listening... Speak continuously. Say 'stop recording' to stop.")

        audio_data = []
        full_text = ""

        while True:
            try:
                # Listen without a timeout, allowing continuous speech
                audio = recognizer.listen(source)
                text = recognizer.recognize_google(audio)

                if "stop recording" in text.lower():
                    st.warning("Voice command detected: Stopping recording.")
                    break  # Stop listening when "stop recording" is detected

                audio_data.append(text)
                full_text = " ".join(audio_data)
                st.session_state.recognized_text = full_text  

                # Update UI dynamically
                st.text_area("Recognized Speech (Live):", full_text, height=150)
                st.rerun()  

            except sr.UnknownValueError:
                st.warning("Couldn't understand the audio. Please speak clearly.")
            except sr.RequestError:
                st.error("Could not connect to Google Speech Recognition. Check internet.")
                break

    return full_text


if st.button("🎙 Start Speech Recognition", key="recognize_btn"):
    result = recognize_speech()
    if result:
        st.success("Speech recognized successfully!")

st.text_area("Recognized Text (Editable):", value=st.session_state.recognized_text, height=150)

st.divider()

### 🌍 Translation Section
st.subheader("🌍 Translation")

def translate_text(text, target_lang):
    """Function to translate text to the selected language"""
    if not text:
        st.warning("Please provide text to translate.")
        return None
    
    with st.spinner("Translating..."):
        try:
            translator = GoogleTranslator(source='auto', target=target_lang)
            translated = translator.translate(text)
            
            # Update session state
            st.session_state.translated_text = translated
            
            # Add to history
            st.session_state.translation_history.append({
                "original": text,
                "translated": translated,
                "language": target_lang
            })
            
            return translated
        except Exception as e:
            st.error(f"Translation error: {str(e)}")
            return None

if st.button("🌍 Translate Text", key="translate_btn"):
    if st.session_state.recognized_text:
        translate_text(st.session_state.recognized_text, st.session_state.target_language)
        st.success("Translation completed!")
    else:
        st.warning("Please recognize speech first.")

st.text_area("Translated Text:", value=st.session_state.translated_text, height=150, disabled=True)

st.divider()

### 💾 Save and Download Section
st.subheader("💾 Save and Download")

def create_document():
    """Function to create a Word document with the recognition and translation results"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Speech Recognition and Translation Results', 0)
    
    # Add current translation if available
    if st.session_state.recognized_text:
        doc.add_heading('Current Translation', level=1)
        doc.add_paragraph(f"Original text: {st.session_state.recognized_text}")
        
        if st.session_state.translated_text:
            target_lang_name = next((name for name, code in language_options.items() 
                                    if code == st.session_state.target_language), st.session_state.target_language)
            doc.add_paragraph(f"Translated to {target_lang_name}: {st.session_state.translated_text}")
    
    # Add history
    if st.session_state.translation_history:
        doc.add_heading('Translation History', level=1)
        
        for i, item in enumerate(st.session_state.translation_history):
            doc.add_heading(f"Translation {i+1}", level=2)
            doc.add_paragraph(f"Original text: {item['original']}")
            
            target_lang_name = next((name for name, code in language_options.items() 
                                    if code == item['language']), item['language'])
            doc.add_paragraph(f"Translated to {target_lang_name}: {item['translated']}")
    
    # Save document to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("📄 Generate Document"):
    st.session_state.doc_buffer = create_document()
    st.success("Document ready! Download below.")

if 'doc_buffer' in st.session_state:
    st.download_button(
        label="📥 Download Word Document",
        data=st.session_state.doc_buffer,
        file_name=f"speech_translation_{time.strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def create_text_file():
    """Function to create a text file with the recognized and translated text."""
    text_content = "Speech Recognition Result:\n" + st.session_state.recognized_text + "\n\n"
    if st.session_state.translated_text:
        text_content += "Translated Text:\n" + st.session_state.translated_text + "\n"
    
    buffer = io.BytesIO()
    buffer.write(text_content.encode())
    buffer.seek(0)
    return buffer

st.download_button(
    label="📜 Download as Text File",
    data=create_text_file(),
    file_name="speech_translation.txt",
    mime="text/plain"
)

st.divider()

### 📖 Translation History Section
with st.expander("📖 View Translation History", expanded=False):
    for i, item in enumerate(st.session_state.translation_history):
        st.write(f"**Translation {i+1}:**")
        st.write(f"Original: {item['original']}")
        st.write(f"Translated to {item['language']}: {item['translated']}")
        st.write("---")
